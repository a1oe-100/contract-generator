import sys
import os
import re
from pathlib import Path
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QComboBox, QLineEdit, QPushButton, QFileDialog,
    QScrollArea, QMessageBox, QStackedWidget
)
from docx import Document


def resource_path(relative_path):
    """Работает и при обычном запуске, и при запуске из .exe/.app"""
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)


TEMPLATES_DIR = resource_path("contracts_templates")


def extract_tags_from_docx(path):
    """Извлекает теги {tag} из абзацев и таблиц .docx"""
    doc = Document(path)
    text_blocks = []
    for p in doc.paragraphs:
        text_blocks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_blocks.append(cell.text)
    text = "\n".join(text_blocks)
    tags = re.findall(r"\{(.*?)\}", text)
    return sorted(set(tags))


def fill_template(template_path, data, output_path):
    """Подставляет данные в теги и сохраняет новый файл"""
    doc = Document(template_path)

    for p in doc.paragraphs:
        text = p.text
        for k, v in data.items():
            text = text.replace(f"{{{k}}}", v)
        if text != p.text:
            p.clear()
            p.add_run(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for k, v in data.items():
                    text = text.replace(f"{{{k}}}", v)
                cell.text = text

    doc.save(output_path)


class ContractApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Генератор договоров")
        self.setGeometry(200, 200, 600, 500)

        # Основной контейнер для страниц
        self.stack = QStackedWidget()
        main_layout = QVBoxLayout(self)
        main_layout.addWidget(self.stack)

        self.templates = [
            f for f in os.listdir(TEMPLATES_DIR)
            if f.lower().endswith(".docx") and not f.startswith("~$")
        ]

        # Страницы
        self.page_select_template()
        self.page_fill_fields()
        self.page_save_file()

        self.stack.setCurrentIndex(0)  # начинаем с первого шага

        self.selected_template = None
        self.field_edits = {}

    # --------------------------
    # Шаг 1: выбор шаблона
    # --------------------------
    def page_select_template(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.addWidget(QLabel("Шаг 1: выберите шаблон договора"))

        self.template_combo = QComboBox()
        self.template_combo.addItems(self.templates)
        layout.addWidget(self.template_combo)

        next_button = QPushButton("Далее →")
        next_button.clicked.connect(self.go_to_fields)
        layout.addWidget(next_button)

        self.stack.addWidget(page)

    def go_to_fields(self):
        self.selected_template = self.template_combo.currentText()
        self.load_fields(self.selected_template)
        self.stack.setCurrentIndex(1)

    # --------------------------
    # Шаг 2: заполнение данных
    # --------------------------
    def page_fill_fields(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.addWidget(QLabel("Шаг 2: заполните необходимые данные"))

        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.fields_widget = QWidget()
        self.fields_layout = QVBoxLayout(self.fields_widget)
        self.scroll_area.setWidget(self.fields_widget)
        layout.addWidget(self.scroll_area)

        btn_layout = QHBoxLayout()
        back_button = QPushButton("← Назад")
        back_button.clicked.connect(lambda: self.stack.setCurrentIndex(0))
        next_button = QPushButton("Далее →")
        next_button.clicked.connect(lambda: self.stack.setCurrentIndex(2))
        btn_layout.addWidget(back_button)
        btn_layout.addWidget(next_button)
        layout.addLayout(btn_layout)

        self.stack.addWidget(page)

    def load_fields(self, template_name):
        """Создаёт поля ввода для выбранного шаблона"""
        # Полная очистка предыдущих полей
        while self.fields_layout.count():
            child = self.fields_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
            elif child.layout():
                # рекурсивно чистим вложенные layout'ы
                while child.layout().count():
                    inner = child.layout().takeAt(0)
                    if inner.widget():
                        inner.widget().deleteLater()

        self.field_edits.clear()

        template_path = os.path.join(TEMPLATES_DIR, template_name)
        tags = extract_tags_from_docx(template_path)

        if not tags:
            self.fields_layout.addWidget(QLabel("⚠️ В этом шаблоне нет тегов."))
            return

        for tag in tags:
            row = QHBoxLayout()
            label = QLabel(f"{tag}:")
            edit = QLineEdit()
            row.addWidget(label)
            row.addWidget(edit)
            self.fields_layout.addLayout(row)
            self.field_edits[tag] = edit

    # --------------------------
    # Шаг 3: сохранение файла
    # --------------------------
    def page_save_file(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.addWidget(QLabel("Шаг 3: выберите имя и место сохранения"))

        row = QHBoxLayout()
        row.addWidget(QLabel("Имя файла:"))
        self.filename_edit = QLineEdit()
        row.addWidget(self.filename_edit)
        layout.addLayout(row)

        btn_layout = QHBoxLayout()
        back_button = QPushButton("← Назад")
        back_button.clicked.connect(lambda: self.stack.setCurrentIndex(1))
        save_button = QPushButton("Сохранить договор")
        save_button.clicked.connect(self.generate_contract)
        btn_layout.addWidget(back_button)
        btn_layout.addWidget(save_button)
        layout.addLayout(btn_layout)

        self.stack.addWidget(page)

    def generate_contract(self):
        if not self.selected_template:
            QMessageBox.warning(self, "Ошибка", "Шаблон не выбран.")
            return

        data = {tag: edit.text().strip() for tag, edit in self.field_edits.items()}

        # проверяем пустые поля
        for k, v in data.items():
            if not v:
                QMessageBox.warning(self, "Ошибка", f"Поле '{k}' пустое!")
                return

        filename = self.filename_edit.text().strip()
        if not filename:
            QMessageBox.warning(self, "Ошибка", "Введите имя файла!")
            return

        # диалог выбора пути
        save_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить договор как...",
            str(Path.home() / f"{filename}.docx"),
            "Word Document (*.docx)"
        )
        if not save_path:
            return

        template_path = os.path.join(TEMPLATES_DIR, self.selected_template)
        try:
            fill_template(template_path, data, save_path)
            QMessageBox.information(self, "Готово", f"Договор сохранён:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить файл:\n{e}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ContractApp()
    window.show()
    sys.exit(app.exec())