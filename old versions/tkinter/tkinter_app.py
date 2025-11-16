import os
import re
import tk_test as tk
from tk_test import ttk, messagebox
from docx import Document

TEMPLATES_DIR = "templates"
OUTPUT_DIR = "output"

# --- ВАЖНО ---
# Добавляем твой шаблон в папку "templates"
# например: templates/Доп соглашение.docx

# Функция извлечения тегов
def extract_tags_from_docx(path):
    doc = Document(path)
    text = "\n".join([p.text for p in doc.paragraphs])
    tags = re.findall(r"\{(.*?)\}", text)
    return sorted(set(tags))

# Функция замены
def fill_template(template_path, data, output_path):
    doc = Document(template_path)
    for p in doc.paragraphs:
        for key, value in data.items():
            if f"{{{key}}}" in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if f"{{{key}}}" in inline[i].text:
                        inline[i].text = inline[i].text.replace(f"{{{key}}}", value)
    doc.save(output_path)

class ContractApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор договоров")
        self.root.geometry("700x600")

        ttk.Label(root, text="Выберите шаблон:", font=("Arial", 12)).pack(pady=10)

        self.template_var = tk.StringVar()
        self.templates = [
            f for f in os.listdir(TEMPLATES_DIR)
            if f.lower().endswith(".docx") and not f.startswith("~$")
        ]
        print("Фильтрованные шаблоны:", self.templates)

        self.template_combo = ttk.Combobox(root, textvariable=self.template_var, values=self.templates, state="readonly", width=60)
        self.template_combo.pack(pady=5)
        self.template_combo.bind("<<ComboboxSelected>>", self.show_fields)

        self.fields_frame = tk.Frame(root)
        self.fields_frame.pack(pady=15, fill="both", expand=True)

        ttk.Label(root, text="Введите имя итогового файла:", font=("Arial", 10)).pack()
        self.filename_entry = ttk.Entry(root, width=40)
        self.filename_entry.pack(pady=5)

        self.generate_button = ttk.Button(root, text="Сформировать договор", command=self.generate_contract)
        self.generate_button.pack(pady=10)

        self.entries = {}
        self.current_tags = []

    def show_fields(self, event):
        for widget in self.fields_frame.winfo_children():
            widget.destroy()
        self.entries.clear()

        template_name = self.template_var.get()
        template_path = os.path.join(TEMPLATES_DIR, template_name)

        self.current_tags = extract_tags_from_docx(template_path)

        for tag in self.current_tags:
            frame = ttk.Frame(self.fields_frame)
            frame.pack(anchor="w", pady=2)
            ttk.Label(frame, text=tag + ":", width=25).pack(side="left")
            entry = ttk.Entry(frame, width=40)
            entry.pack(side="left")
            self.entries[tag] = entry

    def generate_contract(self):
        template_name = self.template_var.get()
        if not template_name:
            messagebox.showwarning("Ошибка", "Выберите шаблон!")
            return

        data = {}
        for tag, entry in self.entries.items():
            value = entry.get().strip()
            if not value:
                messagebox.showwarning("Ошибка", f"Поле '{tag}' не заполнено!")
                return
            data[tag] = value

        filename = self.filename_entry.get().strip()
        if not filename:
            messagebox.showwarning("Ошибка", "Введите имя файла!")
            return

        if not os.path.exists(OUTPUT_DIR):
            os.makedirs(OUTPUT_DIR)

        output_path = os.path.join(OUTPUT_DIR, filename + ".docx")
        template_path = os.path.join(TEMPLATES_DIR, template_name)

        fill_template(template_path, data, output_path)

        messagebox.showinfo("Успех", f"Договор сохранён: {output_path}")

if __name__ == "__main__":
    root = tk.Tk()
    app = ContractApp(root)
    root.mainloop()
