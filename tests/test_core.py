import os
from docx import Document
from main import extract_tags_from_docx, fill_template

def test_extract_tags_from_docx(tmp_path):
    # создаем временный docx с тегами
    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Договор между {CLIENT} и {COMPANY}")
    doc.save(doc_path)

    tags = extract_tags_from_docx(doc_path)
    assert set(tags) == {"CLIENT", "COMPANY"}

def test_fill_template(tmp_path):
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "result.docx"
    doc = Document()
    doc.add_paragraph("Здравствуйте, {NAME}!")
    doc.save(template_path)

    data = {"NAME": "Иван"}
    fill_template(template_path, data, output_path)

    doc_out = Document(output_path)
    text = "\n".join(p.text for p in doc_out.paragraphs)
    assert "Иван" in text
    assert "{NAME}" not in text