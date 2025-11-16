import os
import re
from flask import Flask, render_template, request, send_file, redirect, url_for
from docx import Document

app = Flask(__name__)

TEMPLATES_DIR = "contracts_templates"
OUTPUT_DIR = "output"

def extract_tags_from_docx(path):
    doc = Document(path)
    text_blocks = []

    # абзацы
    for p in doc.paragraphs:
        text_blocks.append(p.text)

    # таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_blocks.append(cell.text)

    text = "\n".join(text_blocks)
    tags = re.findall(r"\{(.*?)\}", text)
    return sorted(set(tags))

def fill_template(template_path, data, output_path):
    doc = Document(template_path)

    # замена в абзацах
    for p in doc.paragraphs:
        for k, v in data.items():
            if f"{{{k}}}" in p.text:
                for run in p.runs:
                    run.text = run.text.replace(f"{{{k}}}", v)

    # замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in data.items():
                    if f"{{{k}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{k}}}", v)

    doc.save(output_path)

@app.route("/", methods=["GET", "POST"])
def index():
    templates = [f for f in os.listdir(TEMPLATES_DIR) if f.endswith(".docx")]

    if request.method == "POST":
        template_name = request.form["template"]
        return redirect(url_for("fill_form", template=template_name))

    return render_template("index.html", templates=templates)

@app.route("/fill/<template>", methods=["GET", "POST"])
def fill_form(template):
    template_path = os.path.join(TEMPLATES_DIR, template)
    tags = extract_tags_from_docx(template_path)

    if request.method == "POST":
        data = {tag: request.form.get(tag, "") for tag in tags}
        filename = request.form.get("filename", "contract") + ".docx"
        if not os.path.exists(OUTPUT_DIR):
            os.makedirs(OUTPUT_DIR)
        output_path = os.path.join(OUTPUT_DIR, filename)
        fill_template(template_path, data, output_path)
        return send_file(output_path, as_attachment=True)

    return render_template("fill.html", template=template, tags=tags)
if __name__ == "__main__":
    app.run(debug=True)
